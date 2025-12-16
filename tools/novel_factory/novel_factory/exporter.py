from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document

from .config import CHAPTERS_DIR, EXPORTS_DIR


def gather_chapter_texts() -> List[tuple[int, str]]:
    texts: List[tuple[int, str]] = []
    for path in sorted(CHAPTERS_DIR.glob("[0-9][0-9][0-9][0-9]")):
        output = path / "output.md"
        if output.exists():
            with output.open("r", encoding="utf-8") as f:
                content = f.read()
            try:
                ch_no = int(path.name)
            except ValueError:
                continue
            texts.append((ch_no, content))
    texts.sort(key=lambda x: x[0])
    return texts


def export_txt(dest: Path | None = None) -> Path:
    dest = dest or EXPORTS_DIR / "novel.txt"
    dest.parent.mkdir(parents=True, exist_ok=True)
    texts = gather_chapter_texts()
    with dest.open("w", encoding="utf-8") as f:
        for ch, txt in texts:
            f.write(f"# Chapter {ch}\n\n")
            f.write(txt.split("===NEXT_CHAPTER_PACK===")[0].replace("===CHAPTER_TEXT===", "").strip())
            f.write("\n\n")
    return dest


def export_docx(dest: Path | None = None) -> Path:
    dest = dest or EXPORTS_DIR / "novel.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for ch, txt in gather_chapter_texts():
        doc.add_heading(f"Chapter {ch}", level=1)
        prose = txt.split("===NEXT_CHAPTER_PACK===")[0].replace("===CHAPTER_TEXT===", "").strip()
        for para in prose.split("\n\n"):
            doc.add_paragraph(para)
    doc.save(dest)
    return dest
